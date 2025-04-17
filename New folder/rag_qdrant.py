from fastapi import FastAPI
import uvicorn
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from PyPDF2 import PdfReader
from langchain_community.document_loaders import PyPDFLoader
import docx
import os
from langchain_groq import ChatGroq
from dotenv import load_dotenv
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain, RetrievalQA
from langchain.memory import ConversationBufferMemory
from langchain_community.callbacks.manager import get_openai_callback
from langchain_core.prompts import PromptTemplate
from langchain.docstore.document import Document
from pinecone import Pinecone
from langchain_pinecone import PineconeVectorStore

pinecone_api_key = "pcsk_4zXG2N_MKp9BDVZhBsknSMSRJRTprmxAFaw69bpz1uQdkeWeLAJU5mCYDAjfMPMbc3C9TJ"
groq_api_key = "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"

os.environ['PINECONE_API_KEY'] = pinecone_api_key
load_dotenv()

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Use ["http://localhost:5173"] for better security
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# pc=Pinecone(api_key=pinecone_api_key)
# pc.create_index(
#         name="medicatbot",
#         dimension=3072,
#         metric="cosine",
#         spec=ServerlessSpec(cloud="aws", region="us-east-1"),
#     )



def get_files_text(uploaded_file):
    text=""
    split_tup = os.path.basename(uploaded_file).split(".")
    file_extension = split_tup[1]
    print(file_extension)
    if file_extension == "pdf":
        text = get_pdf_text(uploaded_file)
    elif file_extension == "docx":
        text = get_docx_text(uploaded_file)
    else:
        pass
    return text

def get_pdf_text(pdf_file):
    text=""
    pdf_reader=PyPDFLoader(pdf_file)
    doc=pdf_reader.load()
    for page in doc:
        text+=page.page_content
    return text

def get_docx_text(docxs):
    all_text = []
    doc = docx.Document(docxs)
    for para in doc.paragraphs:
        all_text.append(para.text)
    text=''.join(all_text)
    return text

def get_csv_text():
    return "a"

def get_text_chunks(text,filename):
    text_splitter = CharacterTextSplitter(
        separator='\n', chunk_size=4096, chunk_overlap=300, length_function=len
    )
    chunks = text_splitter.split_text(text)

    doc_list=[]
    for chunk in chunks:
        metadata = {"source": filename}
        doc_string = Document(page_content=chunk, metadata=metadata)
        doc_list.append(doc_string)

    return doc_list

def get_embeddings():
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return embeddings
# def get_vectorstore(text_chunks):
    
#     vector_store = PineconeVectorStore.from_documents(text_chunks,index_name="medicalbot", embedding=embeddings)
#     # knowledge_base=vector_store.add_documents(text_chunks)
#     return vector_store



def get_vectorstore(embeddings):
    # Using the hugging face embedding models
    # embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    # creating the Vector Store using Facebook AI Semantic search
    index = "medicalbot"
    knowledge_base = PineconeVectorStore.from_existing_index(index, embeddings)

    return knowledge_base



llm = ChatGroq(model="Llama3-70b-8192", groq_api_key=groq_api_key, temperature=0.5)
embeddings = get_embeddings()

knowledge_base = get_vectorstore(embeddings)
print("Knowledge Base Created and data uploaded to vector store")



# uploaded_files = ["First Aid for the USMLE Step 1 2024 34th Edition.pdf"]

# for file  in uploaded_files:
#     file_name = os.path.basename(file)
#     text = get_files_text(file_name)

#     print(file_name + " File Loaded......")
#     # print(text)
#     text_chunks = get_text_chunks(text, file)
#     print("Chunks Created")


question="What is the treatment for diabetes?"

prompt_Template = """
You are a Healthcare assistant specialized in recognizing diseases with their symptoms and also recommend medicine and treatment.
When given a prompt, you will generate output and give the disease's description, symptoms, and treatment.
You are also responsible for general and specific queries according to the healthcare.

Context: {context}
Question: {question}

Only return the helpful answer below. If the question is not related to the context,
politely respond that you are tuned to only answer questions that are related to the context.
"""

PROMPT = PromptTemplate(template=prompt_Template, input_variables=["context", "question"])
chain_type_kwargs = {"prompt": PROMPT}

qa_chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    retriever=knowledge_base.as_retriever(),
    memory=ConversationBufferMemory(memory_key='chat_history', return_messages=True),
    combine_docs_chain_kwargs=chain_type_kwargs
)

rqa = RetrievalQA.from_chain_type(
    llm=llm,
    chain_type="stuff",
    retriever=knowledge_base.as_retriever(search_kwargs={"k": 1}),
    return_source_documents=True,
    chain_type_kwargs=chain_type_kwargs,
)

class MessageRequest(BaseModel):
    message: str

# @app.post("/getresponse")
# def get_response(request: MessageRequest):
#     response = qa_chain({'question': request.message})
#     return response

@app.post("/getresponse")
def get_model_res(request: MessageRequest):
    response = rqa({'query': request.message})
    return response

# qa_res = get_response(question)
# rqa_res = get_model_res(question)

# print("QAC RESPONSE: \n", qa_res)
# print("RQA RESPONSE: \n" ,rqa_res)

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
