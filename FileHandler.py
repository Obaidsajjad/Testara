from langchain_community.document_loaders import PyPDFLoader
import docx
import os
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document
from langchain_community.vectorstores import FAISS
# from ResearchAgent import llm, embeddings
from langchain.chains import RetrievalQA
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
import streamlit as st

llm = ChatGroq(groq_api_key=os.getenv('GROQ_API_KEY'), model_name='Llama3-70b-8192', temperature=0.5)
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")


class FileHandler:
    def get_files_text(self,uploaded_file):
        text=""
        split_tup = os.path.basename(uploaded_file).split(".")
        file_extension = split_tup[1]
        print(file_extension)
        if file_extension == "pdf":
            text = self.get_pdf_text(uploaded_file)
        elif file_extension == "docx":
            text = self.get_docx_text(uploaded_file)
        else:
            pass
        return text

    def get_pdf_text(self,pdf_file):
        text=""
        pdf_reader=PyPDFLoader(pdf_file)
        doc=pdf_reader.load()
        for page in doc:
            text+=page.page_content
        return text

    def get_docx_text(self,docxs):
        all_text = []
        doc = docx.Document(docxs)
        for para in doc.paragraphs:
            all_text.append(para.text)
        text=''.join(all_text)
        return text

    def get_csv_text(self):
        return "a"

    def get_text_chunks(self,text,filename):
        text_splitter = CharacterTextSplitter(
            separator='\n', chunk_size=4096, chunk_overlap=300, length_function=len
        )
        chunks = text_splitter.split_text(text)

        doc_list=[]
        for chunk in chunks:
            metadata = {"source": filename}
            doc_string = Document(page_content=chunk, metadata=metadata)
            doc_list.append(doc_string)

        return doc_list

    def get_knowledgebase(self,chunks,embeddings):
        vectorstore = FAISS.from_documents(chunks, embeddings)
        return vectorstore
    
    def pipeline(self):
        uploaded_files = ["Medical_book.pdf"]
        for file  in uploaded_files:
            file_name = os.path.basename(file)
            text = self.get_files_text(file)

            print(file_name + " File Loaded......")
            # print(text)
            text_chunks = self.get_text_chunks(text, file)
            print("Chunks Created")
            knowledge_base = self.get_knowledgebase(text_chunks,embeddings)
            qa = RetrievalQA.from_chain_type(
                        llm=llm, 
                        chain_type="stuff",
                        retriever=knowledge_base.as_retriever(search_kwargs={"k": 2}),
                        return_source_documents=True,
                        # chain_type_kwargs=chain_type_kwargs,
                    )
            return qa

st.write("FileHandler.py Loaded")
        
print("FileHandler.py loaded")